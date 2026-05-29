"""
Пояснительная записка по курсовой работе.
Тема: «Система управления базами данных, вариант 1 (B-tree)».

Реализовано:
- Пункт 0 (обязательный): базовая файловая СУБД с B-tree индексом.
- Дополнительные задания: 3 (Клиент-Сервер, gRPC), 7 (Access Logs),
  8 (Telemetry), 10 (DEFAULT), 11 (AND/OR в WHERE), 12 (SUM/COUNT/AVG).

Форматирование по ТЗ cw2026:
- Поля: левое 20мм, остальные 15мм.
- Заголовки: Times New Roman 16pt, интервал 1.5, по левому краю.
- Основной текст: Times New Roman 14pt, интервал 1.15, по ширине, отступ 1.25 см.
- Подписи рисунков/таблиц/листингов: TNR 12pt курсив, интервал 1.0.
- Листинги: Consolas 12pt, интервал 1.0.
- Нумерация страниц: с титульного, по центру; на титульном номер не указан.
- Список источников: ГОСТ Р 7.0.100-2018.
"""

import os

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Стили по ТЗ ────────────────────────────────────────────────────────────
FONT_TEXT = "Times New Roman"
FONT_CODE = "Consolas"

SIZE_HEADING = 16
SIZE_TEXT    = 14
SIZE_CAPTION = 12
SIZE_CODE    = 12

SPACING_HEADING = 1.5
SPACING_TEXT    = 1.15
SPACING_CAPTION = 1.0
SPACING_CODE    = 1.0


def set_run_font(run, name=FONT_TEXT, size=SIZE_TEXT, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


def set_paragraph_format(p, spacing, alignment, indent_first=True):
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = alignment
    if indent_first:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_HEADING, WD_ALIGN_PARAGRAPH.LEFT,
                         indent_first=False)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.style = doc.styles[f'Heading {level}']
    # сбрасываем встроенный стиль heading на TNR 16 жирный
    for r in p.runs:
        r.text = ''
    run = p.add_run(text)
    set_run_font(run, FONT_TEXT, SIZE_HEADING, bold=True)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def add_text(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_TEXT, WD_ALIGN_PARAGRAPH.JUSTIFY,
                         indent_first=True)
    run = p.add_run(text)
    set_run_font(run, FONT_TEXT, SIZE_TEXT)
    return p


def add_listing_caption(doc, n, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_CAPTION, WD_ALIGN_PARAGRAPH.LEFT,
                         indent_first=False)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Листинг {n}. {text}")
    set_run_font(run, FONT_TEXT, SIZE_CAPTION, italic=True)
    return p


def add_listing(doc, code):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_CODE, WD_ALIGN_PARAGRAPH.LEFT,
                         indent_first=False)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(code)
    set_run_font(run, FONT_CODE, SIZE_CODE)
    return p


def add_figure_caption(doc, n, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_CAPTION, WD_ALIGN_PARAGRAPH.CENTER,
                         indent_first=False)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Рисунок {n}. {text}")
    set_run_font(run, FONT_TEXT, SIZE_CAPTION, italic=True)
    return p


def add_table_caption(doc, n, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, SPACING_CAPTION, WD_ALIGN_PARAGRAPH.LEFT,
                         indent_first=False)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Таблица {n}. {text}")
    set_run_font(run, FONT_TEXT, SIZE_CAPTION, italic=True)
    return p


def setup_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    set_run_font(run, FONT_TEXT, SIZE_TEXT)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def configure_section(section, different_first=False):
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.right_margin = Mm(15)
    section.left_margin = Mm(20)

    if different_first:
        sectPr = section._sectPr
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)

    setup_page_numbers(section)


# ── Содержание ────────────────────────────────────────────────────────────

def build_intro(doc):
    add_heading(doc, "Введение", level=1)

    add_text(doc,
        "Тема курсовой работы — файловая СУБД на C++. В одном проекте "
        "приходится одновременно работать с лексическим и синтаксическим "
        "разбором, дисковым хранением, индексной структурой и сетевым "
        "слоем. Получается небольшой, но цельный системный проект.")

    add_text(doc,
        "Цель работы — написать СУБД с собственным SQL-подобным языком "
        "запросов и индексом на основе B-дерева (вариант 1 по ТЗ cw2026).")

    add_text(doc,
        "Чтобы добиться цели, пришлось решить несколько задач. Продумать "
        "структуру хранения «система — база — таблица». Написать лексер "
        "и парсер. Спроектировать формат страниц таблицы и записей. "
        "Реализовать B-дерево на диске и встроить его в исполнитель "
        "запросов. Сделать интерактивный и пакетный режимы CLI. И, "
        "наконец, выполнить часть дополнительных заданий из ТЗ.")

    add_text(doc,
        "По обязательному пункту 0 готова базовая СУБД: типы INT и STRING, "
        "DDL-операторы (CREATE/DROP DATABASE, USE, CREATE/DROP TABLE), "
        "DML-операторы (INSERT, SELECT, UPDATE, DELETE), условия WHERE со "
        "сравнениями, BETWEEN и LIKE, ограничения NOT NULL и INDEXED, "
        "вывод результатов в JSON и дисковый B-tree индекс.")

    add_text(doc,
        "Сверх обязательной части сделаны шесть дополнительных заданий: "
        "задание 3 (клиент-серверная архитектура через gRPC), задание 7 "
        "(журнал активности Access Logs), задание 8 (телеметрия: RPS, "
        "среднее время обработки запроса, Error Rate), задание 10 "
        "(модификатор DEFAULT в CREATE TABLE), задание 11 (составные "
        "условия AND/OR со скобками в WHERE) и задание 12 (агрегаты "
        "SUM, COUNT, AVG в SELECT).")

    add_text(doc,
        "В работе использованы C++17, CMake, GoogleTest, gRPC поверх "
        "Protocol Buffers для сетевого слоя, библиотека nlohmann/json для "
        "JSON-вывода и Docker для упаковки серверного и клиентского "
        "приложений. Объём кода около 9500 строк: примерно 3200 строк — "
        "основная реализация (lexer, parser, executor, CLI), 1500 строк — "
        "хранение и индекс, 600 строк — gRPC, журнал и телеметрия, ещё "
        "3200 строк — автоматические тесты. Всего 133 теста, все проходят.")


# ── ОСНОВНАЯ ЧАСТЬ - пункт 0 ──────────────────────────────────────────────

def build_point0(doc):
    add_heading(doc, "1. Пункт 0. Базовая СУБД", level=1)

    add_text(doc,
        "Эта часть закрывает обязательные требования ТЗ: двухуровневая "
        "структура хранения, записи с типами INT и STRING, операторы DDL и "
        "DML, ограничения NOT NULL и INDEXED, B-tree индекс для ускорения "
        "выборки по индексируемым колонкам, валидация запросов, два режима "
        "работы CLI.")

    # ── 1.1 Архитектура ────────────────────────────────────────────────
    add_heading(doc, "1.1. Архитектура и общая схема", level=2)

    add_text(doc,
        "Код разделён на несколько слоёв. CLI отвечает за ввод и вывод. "
        "SqlApi — обёртка над выполнением одной SQL-команды. Lexer "
        "разбирает строку SQL на токены. Parser собирает из них AST. "
        "Executor выполняет команды над каталогом БД. Catalog поддерживает "
        "структуру «система — база — таблица». Storage работает с файлами "
        "таблиц, Index — с файлами индексов.")

    add_text(doc,
        "Обработка запроса проходит так. Строка SQL попадает в Lexer и "
        "разбивается на токены. Parser строит из токенов AST, упакованный "
        "в std::variant над структурами Statement. Executor валидирует "
        "команду, обращается к Catalog и к слоям Storage и Index, "
        "формирует результат. Результат возвращается в CLI или в "
        "gRPC-обработчик. Общая схема — на рисунке 1.")

    add_figure_caption(doc, 1,
        "Общая схема обработки SQL-запроса (см. theory/111_updated.jpg)")

    add_text(doc,
        "Слои тестируются независимо: каталог проверяется на уровне C++ "
        "без CLI, парсер — отдельно от исполнителя, исполнитель — на "
        "готовых AST, CLI — на SQL-скриптах. Под это в проекте собрано "
        "четыре бинарника GoogleTest: dbms_tests, dbms_sql_tests, "
        "dbms_cli_tests и dbms_tz_point0_tests.")

    # ── 1.2 Лексер ─────────────────────────────────────────────────────
    add_heading(doc, "1.2. Лексический анализатор", level=2)

    add_text(doc,
        "Лексер — класс Lexer (src/sql/lexer.cpp). Принимает на вход строку "
        "SQL, возвращает std::vector<Token>. Структура токена показана в "
        "листинге 1.")

    add_listing_caption(doc, 1, "Структура токена и набор ключевых слов")
    add_listing(doc,
        "enum class TokenType {\n"
        "    eof_token, identifier, number, string_literal, keyword, symbol\n"
        "};\n\n"
        "enum class Keyword {\n"
        "    create_kw, drop_kw, use_kw, insert_kw, into_kw, select_kw,\n"
        "    from_kw, where_kw, update_kw, set_kw, delete_kw, as_kw,\n"
        "    between_kw, and_kw, like_kw, value_kw, values_kw, table_kw,\n"
        "    database_kw, int_kw, string_kw, bool_kw, not_kw, null_kw,\n"
        "    indexed_kw, default_kw, or_kw, sum_kw, count_kw, avg_kw\n"
        "};\n\n"
        "struct Token {\n"
        "    TokenType type;\n"
        "    std::string text;\n"
        "    Keyword keyword;\n"
        "    char symbol;\n"
        "};")

    add_text(doc,
        "Алгоритм простой: посимвольное сканирование с классификацией "
        "текущей лексемы. Пробельные символы пропускаются. Первый значащий "
        "символ задаёт ветку: буква или подчёркивание — это идентификатор "
        "либо ключевое слово; цифра — число; двойная кавычка — строковый "
        "литерал. Одиночные символы и составные операторы (==, !=, <=, >=) "
        "обрабатываются отдельно. По требованию ТЗ ключевые слова "
        "регистронезависимы, но смешанный регистр внутри одного слова "
        "запрещён: лексер проверяет, что идентификатор либо весь в нижнем, "
        "либо весь в верхнем регистре. Строковые литералы только в двойных "
        "кавычках, одиночные кавычки отбрасываются.")

    add_text(doc,
        "Из C++ используются std::string и std::vector для текста и "
        "токенов, enum class для типов токенов и категорий ключевых слов, "
        "классификация символов через std::isalpha, std::isdigit и "
        "std::isalnum.")

    # ── 1.3 Парсер ─────────────────────────────────────────────────────
    add_heading(doc, "1.3. Синтаксический анализатор", level=2)

    add_text(doc,
        "Парсер написан методом рекурсивного спуска (src/sql/parser.cpp). "
        "Класс Parser получает поток токенов и возвращает значение типа "
        "Statement. Statement — это std::variant над всеми возможными "
        "формами SQL-команд. Такой подход даёт единый тип AST: в "
        "исполнителе достаточно std::visit или std::get_if, иерархия "
        "виртуальных методов не нужна.")

    add_text(doc,
        "Грамматика построена по требованиям ТЗ. Команды могут занимать "
        "несколько строк, завершаются «;». Реализованные правила (в "
        "нотации, близкой к EBNF): CREATE DATABASE name; DROP DATABASE "
        "name; USE name; CREATE TABLE name (colDef {, colDef}); DROP "
        "TABLE name; INSERT INTO name (cols) VALUE row {, row}; "
        "SELECT (* | item {, item}) FROM name [WHERE cond]; "
        "UPDATE name SET assign {, assign} [WHERE cond]; "
        "DELETE FROM name [WHERE cond].")

    add_text(doc,
        "Условия WHERE парсятся через три уровня приоритета. Самый "
        "низкий — OR, выше — AND, самый высокий — атомарное условие: "
        "сравнение, BETWEEN, LIKE или вложенное выражение в скобках. "
        "Этому естественно отвечают три взаимно рекурсивные функции: "
        "parse_where_or, parse_where_and, parse_where_atom. Узлы AND и "
        "OR не могут лежать прямо в std::variant (тип получился бы "
        "рекурсивным через самого себя), поэтому они обёрнуты в "
        "std::unique_ptr и попадают в WhereCondition как "
        "std::unique_ptr<WhereAnd> и std::unique_ptr<WhereOr> (листинг 2).")

    add_listing_caption(doc, 2,
        "Узлы AST условий WHERE с поддержкой AND, OR, BETWEEN и LIKE")
    add_listing(doc,
        "struct WhereComparison { Operand lhs; ComparisonOp op; Operand rhs; };\n"
        "struct WhereBetween   { Operand value; Operand low; Operand high; };\n"
        "struct WhereLike      { Operand value; Operand pattern; };\n\n"
        "struct WhereAnd; struct WhereOr;\n\n"
        "using WhereCondition = std::variant<\n"
        "    WhereComparison,\n"
        "    WhereBetween,\n"
        "    WhereLike,\n"
        "    std::unique_ptr<WhereAnd>,\n"
        "    std::unique_ptr<WhereOr>\n"
        ">;\n\n"
        "struct WhereAnd { WhereCondition left; WhereCondition right; };\n"
        "struct WhereOr  { WhereCondition left; WhereCondition right; };")

    add_text(doc,
        "Имена баз, таблиц и колонок ограничены правилами ТЗ: латинские "
        "буквы, цифры и подчёркивания, причём первый символ — не цифра. "
        "Проверка идёт в парсере. Если правило нарушено, бросается "
        "std::runtime_error с описанием места ошибки. Это исключение "
        "ловится в SqlApi и превращается в текстовое сообщение "
        "пользователю.")

    # ── 1.4 Executor ───────────────────────────────────────────────────
    add_heading(doc, "1.4. Семантика и исполнение запросов", level=2)

    add_text(doc,
        "Исполнитель собран в классе Executor (src/sql/executor.cpp, около "
        "2200 строк). Точка входа — метод execute(stmt). Через std::visit "
        "он диспетчеризует выполнение на конкретные методы: "
        "execute_create_database, execute_create_table, execute_insert, "
        "execute_select, execute_update, execute_delete и так далее. "
        "Семантика отдельных операторов описана в приложении Б.")

    add_text(doc,
        "Любой DML-оператор проходит общие этапы. Сначала разрешается имя "
        "таблицы (resolve_table_name): либо через активную базу из USE, "
        "либо по явной нотации database.table. Дальше проверяется WHERE "
        "по типам и именам колонок (validate_where_condition). Потом "
        "пробуется построить индексный план (try_build_index_where_plan). "
        "Записи собираются через индекс или полным сканом. К ним "
        "применяется проекция или модификация, после чего формируется "
        "JSON-результат.")

    add_text(doc,
        "Для INSERT, UPDATE и DELETE поддерживается согласованность файла "
        "таблицы с индексами. Порядок такой. Сначала валидация: типы, NOT "
        "NULL, уникальность INDEXED. Потом запись в .tbl. И только после "
        "удачной записи — синхронные изменения в .idx. Индексы меняются "
        "через apply_index_mutations_with_rollback: если очередная "
        "операция упала на ошибке ввода-вывода, уже применённые изменения "
        "откатываются в обратном порядке, а заголовок таблицы "
        "восстанавливается из ранее прочитанной копии страницы 0. Это "
        "даёт согласованность при штатных сбоях записи без полноценного "
        "WAL.")

    add_text(doc,
        "Каждый оператор возвращает bool — признак успеха. На "
        "семантические ошибки (нет такой базы, нет такой колонки, "
        "неверный тип, нарушение NOT NULL или INDEXED, дубль уникального "
        "ключа) бросается исключение с описанием. Его ловит SqlApi, и "
        "программа не падает. По ТЗ это обязательное требование.")

    # ── 1.5 Хранение ───────────────────────────────────────────────────
    add_heading(doc, "1.5. Хранение данных на диске", level=2)

    add_text(doc,
        "База — это отдельная директория внутри корневого каталога данных. "
        "Путь задаётся переменной окружения DBMS_DATA_ROOT. Каждая таблица "
        "лежит в файле с расширением .tbl, каждый индекс — в файле .idx с "
        "именем вида table__column.idx. Так удобно: DROP TABLE — это "
        "удаление нескольких файлов, а перенос базы — обычное копирование "
        "директории.")

    add_text(doc,
        "Файл .tbl разбит на страницы по 4096 байт. Страница 0 хранит "
        "TableHeader и сериализованную схему колонок: имя, тип, флаги NOT "
        "NULL и INDEXED, значение DEFAULT. Страницы 1..next_page_id-1 "
        "хранят сами записи в slotted-формате (листинг 3). Физический "
        "адрес записи — пара (page_id, slot_id), структура Rid. Именно "
        "Rid лежит в индексе как значение, поэтому индекс не дублирует "
        "сами записи (требование ТЗ).")

    add_listing_caption(doc, 3, "Заголовок страницы и слот записи")
    add_listing(doc,
        "static constexpr std::size_t kPageSize = 4096;\n\n"
        "struct PageHeader {\n"
        "    int slots_count;\n"
        "    int free_start;\n"
        "    int free_end;\n"
        "};\n\n"
        "struct Slot {\n"
        "    int offset;\n"
        "    int size;\n"
        "};\n\n"
        "struct RecordHeader {\n"
        "    int size;\n"
        "    int null_bytes;\n"
        "};")

    add_text(doc,
        "Запись сериализуется функцией serialize_record. Формат такой: "
        "сначала RecordHeader с размером записи и размером null-карты, "
        "потом сама null-карта (битовый вектор, где единичный бит на "
        "позиции i означает NULL в i-й колонке), потом по очереди "
        "значения не-NULL колонок. Числа пишутся как 4-байтовый int, "
        "строки — как (длина: int) + (байты строки). Обратная операция "
        "deserialize_record читает заголовок, восстанавливает null-карту "
        "и складывает значения в Record::values типа "
        "std::vector<std::optional<std::string>>. У формата два плюса: "
        "строки переменной длины не требуют максимального буфера, а NULL "
        "обходится без зарезервированного значения.")

    add_text(doc,
        "Удаление логическое: размер слота ставится в 0, остальные слоты "
        "не сдвигаются. Так не ломаются ранее выданные Rid, в том числе "
        "те, что лежат в индексах по другим колонкам этой же таблицы. "
        "Уплотнения страниц в текущей версии нет, это намеренное "
        "упрощение для пункта 0.")

    # ── 1.6 B-tree ─────────────────────────────────────────────────────
    add_heading(doc, "1.6. B-tree индекс на диске", level=2)

    add_text(doc,
        "Индекс — шаблонный класс BTreeDiskIndex<tkey, t> "
        "(include/dbms/index/b_tree_disk_index.h, около 570 строк). "
        "Параметр t — степень дерева, в работе используется t = 5: "
        "минимум 4 ключа в узле, максимум 9. Шаблонность по tkey "
        "закрывает оба случая из ТЗ — INT INDEXED и STRING INDEXED. Для "
        "строк определена обёртка StringKey фиксированного размера, "
        "которая удовлетворяет условию "
        "std::is_trivially_copyable_v<tkey>.")

    add_text(doc,
        "Раскладка индексного файла повторяет идею файла таблицы. "
        "Страница 0 — это IndexHeader: идентификатор корня, порядок "
        "дерева и следующий свободный id узла. Страницы 1..next_page_id-1 "
        "— сами узлы. В узле хранится: флаг is_leaf, текущее число "
        "ключей, массив ключей размера 2t−1, массив значений типа Rid "
        "размера 2t−1 и массив идентификаторов потомков размера 2t "
        "(листинг 4).")

    add_listing_caption(doc, 4, "Узел B-дерева на диске")
    add_listing(doc,
        "template<typename tkey, std::size_t t = 5>\n"
        "class BTreeDiskIndex {\n"
        "    static constexpr int kMinKeys     = static_cast<int>(t - 1);\n"
        "    static constexpr int kMaxKeys     = static_cast<int>(2 * t - 1);\n"
        "    static constexpr int kMaxChildren = static_cast<int>(2 * t);\n\n"
        "    struct Node {\n"
        "        int id;\n"
        "        int is_leaf;\n"
        "        int keys_count;\n"
        "        std::array<tkey, kMaxKeys + 1>     keys;\n"
        "        std::array<Rid,  kMaxKeys + 1>     values;\n"
        "        std::array<int,  kMaxChildren + 1> children;\n"
        "    };\n"
        "    // ... insert, find, erase, range\n"
        "};")

    add_text(doc,
        "Реализованы стандартные операции из учебника Кормена. Вставка "
        "(insert) — с проактивным расщеплением полных узлов на пути вниз. "
        "Поиск (find и contains) — обычный спуск по дереву. Удаление "
        "(erase) — с тремя случаями: удаление из листа, слияние потомков "
        "и заимствование ключа у соседнего ребёнка. Интервальный обход "
        "(range) используется для BETWEEN и неравенств. Все операции "
        "обращаются к диску через IndexPageManager: он читает и пишет "
        "страницы целиком и выделяет новые идентификаторы узлов в "
        "IndexHeader.")

    add_text(doc,
        "Индекс используется для оптимизации условий вида indexed_column "
        "op literal, где op — это ==, !=, <, <=, >, >= или BETWEEN "
        "(литерал может стоять и слева). Для каждого подходящего условия "
        "исполнитель выбирает операцию B-дерева (find или range), "
        "получает множество Rid, а потом читает по этим Rid сами записи "
        "из файла таблицы. Условия, которые не сводятся к этому шаблону "
        "(сравнение двух колонок, LIKE по регулярному выражению, сложные "
        "AND/OR), обрабатываются полным сканом страниц. Корректность "
        "выбора пути проверена тестами, в том числе через диагностический "
        "флаг Executor::last_operation_used_index().")

    # ── 1.7 CLI ────────────────────────────────────────────────────────
    add_heading(doc, "1.7. Командный интерфейс", level=2)

    add_text(doc,
        "Командный интерфейс — класс Cli (src/sql/cli.cpp). Режим "
        "переключается аргументом командной строки. Без аргументов "
        "(./prog) запускается интерактивный режим. Пользователь вводит "
        "SQL построчно, Cli накапливает буфер до символа «;», передаёт "
        "команду в SqlApi и печатает результат: короткое OK для DDL/DML "
        "или JSON-массив для SELECT. Приглашение зависит от состояния "
        "буфера: «> » для новой команды, «... » для продолжения "
        "многострочной.")

    add_text(doc,
        "С именем файла (./prog script.sql) включается пакетный режим. "
        "Файл читается целиком, разбивается на команды по «;» вне "
        "строковых литералов и выполняется по порядку. Ошибка одной "
        "команды не валит весь скрипт: сообщение об ошибке печатается, "
        "следующая команда продолжает работу. Этим выполняется требование "
        "ТЗ «программа не должна завершаться аварийно».")


# ── 2. Доп. задание 3. gRPC ───────────────────────────────────────────────

def build_addon_grpc(doc):
    add_heading(doc, "2. Дополнительное задание 3. Клиент-Сервер (gRPC)", level=1)

    add_text(doc,
        "В качестве сетевого слоя выбран gRPC поверх Protocol Buffers v3. "
        "У него два удобства. Во-первых, по одной schema-схеме генерируются "
        "и клиент, и сервер на C++. Во-вторых, транспорт идёт «из коробки», "
        "не нужно руками сериализовать и десериализовать сообщения поверх "
        "сокетов.")

    add_text(doc,
        "Схема сервиса описана в proto/sql_service.proto (листинг 5). API "
        "простой, всего три метода. OpenSession создаёт серверную сессию и "
        "возвращает её строковый идентификатор. Execute принимает "
        "идентификатор сессии и текст SQL, возвращает признак успеха, "
        "признак того, что это SELECT, и либо JSON-результат, либо текст "
        "ошибки. CloseSession освобождает ресурсы сессии. Сессии хранятся "
        "в std::unordered_map под std::mutex, потому что к нему ходят "
        "разные потоки gRPC.")

    add_listing_caption(doc, 5, "Описание сервиса в Protocol Buffers")
    add_listing(doc,
        "syntax = \"proto3\";\n"
        "package dbms.rpc;\n\n"
        "message OpenSessionResponse  { bool ok=1; string session_id=2; string error=3; }\n"
        "message ExecuteRequest       { string session_id=1; string sql=2; }\n"
        "message ExecuteResponse      { bool ok=1; bool is_select=2;\n"
        "                                string json=3; string error=4; }\n"
        "message CloseSessionRequest  { string session_id=1; }\n\n"
        "service SqlService {\n"
        "  rpc OpenSession (OpenSessionRequest)  returns (OpenSessionResponse);\n"
        "  rpc Execute     (ExecuteRequest)       returns (ExecuteResponse);\n"
        "  rpc CloseSession(CloseSessionRequest)  returns (CloseSessionResponse);\n"
        "}")

    add_text(doc,
        "Сервер (src/grpc/grpc_server_main.cpp) принимает аргументы "
        "«host:port» и путь к каталогу данных. Внутри обработчика Execute "
        "вызывается тот же путь, что и в локальном CLI: SqlApi → Lexer → "
        "Parser → Executor. Поведение получается одинаковым в локальном и "
        "сетевом режимах, и не приходится дублировать код.")

    add_text(doc,
        "Клиент (src/grpc/grpc_client_main.cpp) умеет оба режима, "
        "интерактивный и пакетный, и по поведению повторяет локальный CLI: "
        "многострочные команды до «;», разные приглашения, обработка "
        "нескольких команд в одной строке. Идея — чтобы пользователь не "
        "замечал разницы между локальной СУБД и сервером по сети.")

    add_text(doc,
        "Сервер и клиент собраны в Docker-образы awak3r/dbms-server и "
        "awak3r/dbms-client и опубликованы на Docker Hub. Сервер монтирует "
        "том /data под базы и слушает порт 50051.")


# ── 3. Доп. задание 7. Access Logs ────────────────────────────────────────

def build_addon_logs(doc):
    add_heading(doc, "3. Дополнительное задание 7. Логирование активности", level=1)

    add_text(doc,
        "Подсистема логов — класс AccessLogger "
        "(include/dbms/grpc/access_logger.h, src/grpc/access_logger.cpp), "
        "подключённый к gRPC-обработчику Execute. На каждый запрос в "
        "журнал пишется строка с фиксированным набором полей: "
        "идентификатор клиента (берётся из peer-адреса gRPC), "
        "идентификатор обработчика (имя RPC-метода), тело запроса (текст "
        "SQL), время начала и время конца обработки, признак успеха и "
        "статус (OK или текст исключения).")

    add_listing_caption(doc, 6, "Интерфейс AccessLogger")
    add_listing(doc,
        "class AccessLogger {\n"
        "public:\n"
        "    explicit AccessLogger(const std::filesystem::path& log_path);\n\n"
        "    void log(const std::string& client_id,\n"
        "             const std::string& handler_id,\n"
        "             const std::string& body,\n"
        "             std::chrono::system_clock::time_point start,\n"
        "             std::chrono::system_clock::time_point end,\n"
        "             bool ok,\n"
        "             const std::string& status);\n\n"
        "private:\n"
        "    std::ofstream _file;\n"
        "    std::mutex    _mutex;\n"
        "};")

    add_text(doc,
        "Файл открывается в режиме append. Одновременная запись из разных "
        "gRPC-потоков защищена std::mutex, поэтому строки в файле не "
        "перемешиваются. Время пишется в формате ISO 8601 с миллисекундной "
        "точностью через std::chrono::system_clock и std::put_time. Этот "
        "формат и читается глазами, и легко разбирается машинными "
        "анализаторами.")


# ── 4. Доп. задание 8. Telemetry ──────────────────────────────────────────

def build_addon_telemetry(doc):
    add_heading(doc, "4. Дополнительное задание 8. Подсистема телеметрии", level=1)

    add_text(doc,
        "Подсистема телеметрии — класс TelemetryCollector "
        "(include/dbms/grpc/telemetry.h, src/grpc/telemetry.cpp). Он "
        "считает в реальном времени четыре метрики из ТЗ: текущий RPS "
        "(число запросов в секунду), средний и максимальный RPS за "
        "последние 10 минут, среднее время обработки запроса за последние "
        "10 секунд и Error Rate за последнюю минуту.")

    add_text(doc,
        "В основе — скользящее окно. Каждой секунде соответствует один "
        "элемент кольцевого массива длиной kWindowMax = 600 (10 минут — "
        "это максимальное окно из ТЗ). Элемент — структура Bucket: "
        "Unix-секунда, число успешных запросов, число ошибок и суммарное "
        "время обработки в микросекундах (листинг 7). На каждый вызов "
        "record() из gRPC-обработчика счётчики соответствующего бакета "
        "увеличиваются. Когда наступает новая секунда, бакет "
        "переинициализируется новым timestamp и нулевыми счётчиками. Так "
        "старые значения автоматически «истекают», и не приходится "
        "вручную чистить память.")

    add_listing_caption(doc, 7,
        "Структура бакета телеметрии и кольцевого окна")
    add_listing(doc,
        "class TelemetryCollector {\n"
        "    struct Bucket {\n"
        "        int64_t  timestamp = 0;\n"
        "        uint64_t requests  = 0;\n"
        "        uint64_t errors    = 0;\n"
        "        uint64_t total_us  = 0;\n"
        "    };\n\n"
        "    static constexpr int kWindowMax = 600; // 10 минут\n\n"
        "    std::array<Bucket, kWindowMax> _buckets{};\n"
        "    std::mutex                     _mutex;\n"
        "    std::ofstream                  _file;\n"
        "    std::thread                    _thread;\n"
        "    std::atomic<bool>              _stop{false};\n"
        "};")

    add_text(doc,
        "Подсчёт и вывод метрик идут в фоновом потоке (background_loop). "
        "Поток просыпается раз в секунду, суммирует значения по нужным "
        "бакетам и пишет строку отчёта. Поток завершается через "
        "std::atomic<bool> _stop, который выставляется в деструкторе. Для "
        "синхронизации записи и чтения бакетов используется std::mutex. "
        "Версия на atomic-полях усложнила бы код без заметного выигрыша по "
        "производительности.")


# ── 5. Доп. задание 10. DEFAULT ───────────────────────────────────────────

def build_addon_default(doc):
    add_heading(doc, "5. Дополнительное задание 10. Модификатор DEFAULT", level=1)

    add_text(doc,
        "DEFAULT добавлен как расширение грамматики CREATE TABLE: после "
        "типа колонки можно указать DEFAULT <литерал>. Значение пишется в "
        "структуру ColumnDef в поле std::optional<std::string> "
        "default_value и попадает в сериализованную схему таблицы рядом с "
        "типом и флагами.")

    add_text(doc,
        "При INSERT для каждой пропущенной в списке колонки проверяется "
        "default_value. Если оно есть — записывается как значение поля. "
        "Если default_value нет, а колонка не помечена NOT NULL — пишется "
        "NULL. Если колонка NOT NULL и DEFAULT не задан — бросается "
        "ошибка с описанием. Поведение совпадает с тем, как DEFAULT "
        "работает в обычных промышленных СУБД, и соответствует ТЗ.")

    add_text(doc,
        "Для INDEXED-колонок DEFAULT тоже разрешён, но он не может быть "
        "NULL: INDEXED запрещает NULL по ТЗ. Эта проверка идёт на стадии "
        "CREATE TABLE; при нарушении таблица не создаётся.")


# ── 6. Доп. задание 11. WHERE AND/OR ──────────────────────────────────────

def build_addon_andor(doc):
    add_heading(doc, "6. Дополнительное задание 11. Составные условия в WHERE", level=1)

    add_text(doc,
        "Расширение WHERE сделано через два новых узла AST — WhereAnd и "
        "WhereOr — и три взаимно рекурсивные функции парсера: "
        "parse_where_or, parse_where_and, parse_where_atom. "
        "parse_where_or принимает любое число условий, разделённых OR. "
        "parse_where_and — любое число условий, разделённых AND. "
        "parse_where_atom распознаёт одно элементарное условие "
        "(сравнение, BETWEEN, LIKE) или вложенное выражение в круглых "
        "скобках. Такая структура автоматически даёт AND приоритет выше "
        "OR — как в стандарте SQL и в обычных промышленных СУБД.")

    add_text(doc,
        "Узлы хранят детей по значению в виде WhereCondition. Из-за этого "
        "тип получался бы рекурсивным через самого себя, поэтому WhereAnd "
        "и WhereOr пришлось обернуть в std::unique_ptr. Это стандартный "
        "приём для рекурсивных AST на std::variant в C++.")

    add_text(doc,
        "Исполнитель обходит дерево условий рекурсивно. Для составных "
        "узлов вычисляются значения левого и правого поддерева и "
        "комбинируются логической операцией. Для атомарных условий "
        "вызывается evaluate_where_condition. Использование индекса для "
        "составных условий не реализовано: оптимизация индексных планов "
        "под AND/OR в общем случае требует отдельной фазы планирования и "
        "в обязательную часть не входила.")


# ── 7. Доп. задание 12. SUM/COUNT/AVG ─────────────────────────────────────

def build_addon_aggregates(doc):
    add_heading(doc, "7. Дополнительное задание 12. Агрегатные функции", level=1)

    add_text(doc,
        "Агрегаты в SELECT поддержаны в формате FUNC(column_name), где "
        "FUNC — это SUM, COUNT или AVG. В AST элемент списка проекций — "
        "структура SelectItem с опциональным полем "
        "std::optional<AggregateFunc> aggregate. Если поле пустое "
        "(std::nullopt) — колонка выводится как обычно. Если заполнено — "
        "исполнитель переключается в агрегатный режим.")

    add_text(doc,
        "Алгоритм агрегации простой. Исполнитель собирает подходящие "
        "строки тем же путём, что и при обычном SELECT (через индекс или "
        "полный скан), но вместо JSON-массива со всеми строками за один "
        "проход накапливает по каждому агрегатному выражению счётчик и "
        "сумму. COUNT учитывает только не-NULL значения. SUM суммирует "
        "числа (NULL пропускается). AVG возвращает SUM/COUNT с двумя "
        "знаками после запятой. Для пустого результата AVG возвращает "
        "null. В JSON выводится одна агрегатная строка, а не множество "
        "исходных.")


# ── Вывод ─────────────────────────────────────────────────────────────────

def build_conclusion(doc):
    add_heading(doc, "Вывод", level=1)

    add_text(doc,
        "В ходе работы спроектирована и реализована файловая СУБД на "
        "C++17. Обязательный пункт 0 закрыт полностью. Поддерживаются "
        "DDL-операторы (CREATE/DROP DATABASE, USE, CREATE/DROP TABLE) и "
        "DML-операторы (INSERT, SELECT, UPDATE, DELETE), типы INT и "
        "STRING, ограничения NOT NULL и INDEXED, условия WHERE со "
        "сравнениями, BETWEEN и LIKE, вывод результатов в JSON, дисковый "
        "B-tree индекс.")

    add_text(doc,
        "Помимо обязательной части выполнено шесть дополнительных заданий "
        "из ТЗ: клиент-серверная архитектура на gRPC и Protocol Buffers "
        "(№ 3), журнал активности (№ 7), телеметрия со скользящим окном "
        "(№ 8), модификатор DEFAULT в CREATE TABLE (№ 10), составные "
        "условия AND и OR со скобками в WHERE (№ 11), агрегаты SUM, COUNT "
        "и AVG в SELECT (№ 12).")

    add_text(doc,
        "Корректность проверена 133 автоматическими тестами GoogleTest в "
        "четырёх бинарниках. Тесты покрывают лексер, парсер, исполнитель, "
        "индекс, сценарии CLI и сценарии пункта 0 из ТЗ. Полный прогон "
        "всех тестов идёт около 100 миллисекунд. К защите также "
        "подготовлены демонстрационные SQL-скрипты в каталоге scripts/: "
        "demo_point0, demo_constraints_errors, demo_restart_seed/check, "
        "demo_grpc.")

    add_text(doc,
        "У текущей версии есть осознанные ограничения. Нет журнала "
        "упреждающей записи (WAL), поэтому восстановление после "
        "аварийного завершения процесса между записями страниц не "
        "гарантируется. Нет ACID-транзакций между несколькими "
        "SQL-командами. Используется модель одного писателя, "
        "конкурентные писатели не поддерживаются. Удаление записей "
        "логическое, страницы не уплотняются. Все эти ограничения "
        "известны и зафиксированы в документации; их можно закрыть в "
        "следующих итерациях проекта.")


# ── Источники ─────────────────────────────────────────────────────────────

def build_sources(doc):
    add_heading(doc, "Список использованных источников", level=1)

    sources = [
        "Силберштац, А. Основы баз данных / А. Силберштац, Г. Корт, С. Сударшан ; "
        "пер. с англ. — 6-е изд. — Москва : Диалектика, 2021. — 1072 с. — "
        "ISBN 978-5-907458-08-8.",

        "Петров, А. Распределённые данные: внутреннее устройство баз данных / "
        "А. Петров ; пер. с англ. — Москва : БХВ-Петербург, 2021. — 368 с. — "
        "ISBN 978-5-9775-6722-4.",

        "gRPC Authors. gRPC Documentation [Электронный ресурс]. — "
        "URL: https://grpc.io/docs/ (дата обращения: 29.05.2026).",

        "Ломанн, Н. JSON for Modern C++ [Электронный ресурс] / Н. Ломанн. — "
        "URL: https://github.com/nlohmann/json (дата обращения: 29.05.2026).",
    ]

    for i, text in enumerate(sources, start=1):
        p = doc.add_paragraph()
        set_paragraph_format(p, SPACING_TEXT, WD_ALIGN_PARAGRAPH.JUSTIFY,
                             indent_first=False)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        run = p.add_run(f"{i}. {text}")
        set_run_font(run, FONT_TEXT, SIZE_TEXT)


# ── Приложения ────────────────────────────────────────────────────────────

def build_appendices(doc):
    add_heading(doc, "Приложения", level=1)

    add_heading(doc, "Приложение А. Структурная схема СУБД", level=2)

    img_path = os.path.join(os.path.dirname(__file__), "..", "111_updated.jpg")
    if os.path.isfile(img_path):
        p = doc.add_paragraph()
        set_paragraph_format(p, 1.0, WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        run = p.add_run()
        run.add_picture(img_path, width=Mm(160))
    add_figure_caption(doc, 2,
        "Общая схема взаимодействия слоёв СУБД")

    add_heading(doc, "Приложение Б. Соответствие пунктов ТЗ и файлов кода",
                level=2)

    items = [
        ("Пункт 0. Лексер",                "src/sql/lexer.cpp, include/dbms/sql/lexer.h"),
        ("Пункт 0. Парсер",                "src/sql/parser.cpp, include/dbms/sql/parser.h, statements.h"),
        ("Пункт 0. Исполнитель и SQL",     "src/sql/executor.cpp, src/sql/sql_api.cpp"),
        ("Пункт 0. Каталог и сущности",    "include/dbms/core/{catalog,database,table,schema,dbms}.h"),
        ("Пункт 0. Хранение таблицы",      "include/dbms/storage/{page,record,record_codec,table_page_manager}.h"),
        ("Пункт 0. B-tree индекс",         "include/dbms/index/{b_tree_disk_index,index_manager,index_page_manager}.h"),
        ("Пункт 0. CLI",                   "src/sql/cli.cpp, src/main.cpp"),
        ("Доп. задание 3. gRPC",           "proto/sql_service.proto, src/grpc/*.cpp"),
        ("Доп. задание 7. Access Logs",    "src/grpc/access_logger.cpp, include/dbms/grpc/access_logger.h"),
        ("Доп. задание 8. Telemetry",      "src/grpc/telemetry.cpp, include/dbms/grpc/telemetry.h"),
        ("Доп. задание 10. DEFAULT",       "src/sql/parser.cpp (parse_column_def), src/sql/executor.cpp"),
        ("Доп. задание 11. AND/OR",        "src/sql/parser.cpp (parse_where_or/and/atom), statements.h"),
        ("Доп. задание 12. SUM/COUNT/AVG", "src/sql/parser.cpp (parse_select_item), src/sql/executor.cpp"),
        ("Тесты",                          "tests/sql/*.cpp, tests/spec/all_tests.cpp, tests/b_tree_tests.cpp"),
        ("Docker",                         "Dockerfile, scripts/demo_*.sql"),
    ]

    add_table_caption(doc, 1, "Соответствие пунктов ТЗ и исходных файлов")
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Пункт ТЗ", "Файлы реализации")):
        cell.text = ''
        p = cell.paragraphs[0]
        set_paragraph_format(p, SPACING_CAPTION, WD_ALIGN_PARAGRAPH.LEFT,
                             indent_first=False)
        run = p.add_run(text)
        set_run_font(run, FONT_TEXT, SIZE_CAPTION, bold=True)
    for i, (a, b) in enumerate(items, start=1):
        row = table.rows[i].cells
        for cell, text in zip(row, (a, b)):
            cell.text = ''
            p = cell.paragraphs[0]
            set_paragraph_format(p, SPACING_CAPTION, WD_ALIGN_PARAGRAPH.LEFT,
                                 indent_first=False)
            run = p.add_run(text)
            set_run_font(run, FONT_TEXT, SIZE_CAPTION)

    add_heading(doc, "Приложение В. Примеры работы программы", level=2)

    add_text(doc,
        "В данном приложении приведён вывод программы при выполнении "
        "демонстрационных SQL-скриптов из каталога scripts/. Скрипты "
        "запускаются командой: DBMS_DATA_ROOT=/tmp/demo ./build/prog scripts/demo_point0.sql")

    # ── В.1 Happy path ────────────────────────────────────────────────
    add_heading(doc, "В.1. Основные операции DDL и DML (demo_point0.sql)", level=3)

    add_listing_caption(doc, 8, "Команды demo_point0.sql")
    add_listing(doc,
        "CREATE DATABASE demo_main;\n"
        "USE demo_main;\n"
        "CREATE TABLE users (id INT INDEXED, login STRING INDEXED,\n"
        "                    age INT NOT NULL, city STRING);\n"
        "INSERT INTO users (id, login, age, city) VALUE\n"
        "  (1, \"alice\", 20, \"Amsterdam\"),\n"
        "  (2, \"bob\",   25, \"Paris\"),\n"
        "  (3, \"carol\", 30, NULL),\n"
        "  (4, \"anna\",  28, \"Rome\");\n"
        "SELECT * FROM users;\n"
        "SELECT id AS user_id, login, age FROM users WHERE id BETWEEN 2 AND 4;\n"
        "SELECT login, city FROM users WHERE login LIKE \"a.*\";\n"
        "UPDATE users SET city = \"Berlin\" WHERE id == 2;\n"
        "DELETE FROM users WHERE login == \"carol\";\n"
        "SELECT login FROM demo_main.users WHERE id == 1;")

    add_listing_caption(doc, 9, "Вывод demo_point0.sql")
    add_listing(doc,
        "OK\nOK\nOK\nOK\n"
        "[\n"
        "  {\"age\": 20, \"city\": \"Amsterdam\", \"id\": 1, \"login\": \"alice\"},\n"
        "  {\"age\": 25, \"city\": \"Paris\",     \"id\": 2, \"login\": \"bob\"},\n"
        "  {\"age\": 30, \"city\": null,          \"id\": 3, \"login\": \"carol\"},\n"
        "  {\"age\": 28, \"city\": \"Rome\",       \"id\": 4, \"login\": \"anna\"}\n"
        "]\n"
        "[\n"
        "  {\"age\": 25, \"login\": \"bob\",   \"user_id\": 2},\n"
        "  {\"age\": 30, \"login\": \"carol\", \"user_id\": 3},\n"
        "  {\"age\": 28, \"login\": \"anna\",  \"user_id\": 4}\n"
        "]\n"
        "[\n"
        "  {\"city\": \"Amsterdam\", \"login\": \"alice\"},\n"
        "  {\"city\": \"Rome\",      \"login\": \"anna\"}\n"
        "]\n"
        "OK\n"
        "[{\"city\": \"Berlin\", \"id\": 2, \"login\": \"bob\"}]\n"
        "OK\n"
        "[\n"
        "  {\"id\": 1, \"login\": \"alice\"},\n"
        "  {\"id\": 2, \"login\": \"bob\"},\n"
        "  {\"id\": 4, \"login\": \"anna\"}\n"
        "]\n"
        "[{\"login\": \"alice\"}]")

    # ── В.2 Ограничения и ошибки ──────────────────────────────────────
    add_heading(doc, "В.2. Ограничения целостности и информативные ошибки (demo_constraints_errors.sql)", level=3)

    add_listing_caption(doc, 10, "Вывод demo_constraints_errors.sql")
    add_listing(doc,
        "OK\nOK\nOK\nOK\n"
        "ERROR: Constraint error: duplicate INDEXED key `1` for column `id`\n"
        "ERROR: Constraint error: duplicate INDEXED key `alice` for column `login`\n"
        "ERROR: Constraint error: column `id` cannot be NULL because it is INDEXED\n"
        "ERROR: Constraint error: required column `age` was omitted from INSERT into `users`\n"
        "ERROR: Type error: value `bad_int` for column `id` is not INT\n"
        "ERROR: Runtime error: invalid LIKE regex `[`: Unexpected character in bracket expression.\n"
        "ERROR: Semantic error: table `demo_errors.users` has no column `missing_column` in SELECT projection\n"
        "[{\"age\": 20, \"id\": 1, \"login\": \"alice\"}]")

    add_text(doc,
        "Все ошибки обрабатываются без аварийного завершения программы. "
        "После серии ошибочных команд финальный SELECT возвращает "
        "корректный результат — таблица осталась в согласованном состоянии.")

    # ── В.3 Персистентность ───────────────────────────────────────────
    add_heading(doc, "В.3. Персистентность данных после перезапуска (demo_restart_seed/check.sql)", level=3)

    add_listing_caption(doc, 11, "Первый запуск — запись данных (demo_restart_seed.sql)")
    add_listing(doc,
        "OK\nOK\nOK\nOK\n"
        "[\n"
        "  {\"balance\": 100, \"id\": 1, \"owner\": \"alice\"},\n"
        "  {\"balance\": 150, \"id\": 2, \"owner\": \"bob\"},\n"
        "  {\"balance\": 200, \"id\": 3, \"owner\": \"carol\"}\n"
        "]")

    add_listing_caption(doc, 12, "Второй запуск — чтение после перезапуска (demo_restart_check.sql)")
    add_listing(doc,
        "OK\n"
        "[\n"
        "  {\"balance\": 100, \"owner\": \"alice\"},\n"
        "  {\"balance\": 150, \"owner\": \"bob\"}\n"
        "]\n"
        "[{\"id\": 3}]")

    add_text(doc,
        "Данные, схемы и индексы сохраняются в файловой системе и "
        "восстанавливаются при повторном запуске на том же DBMS_DATA_ROOT "
        "без дополнительных действий.")


# ── Сборка документа ──────────────────────────────────────────────────────

def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_TEXT
    style.font.size = Pt(SIZE_TEXT)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_TEXT)
    rFonts.set(qn('w:hAnsi'), FONT_TEXT)
    rFonts.set(qn('w:cs'), FONT_TEXT)
    rFonts.set(qn('w:eastAsia'), FONT_TEXT)

    section = doc.sections[0]
    configure_section(section, different_first=True)

    # ── ТИТУЛЬНЫЙ ЛИСТ ────────────────────────────────────────────────
    def title_line(text, bold=False, size=14, before=0, after=0, align='center'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        run = p.add_run(text)
        set_run_font(run, FONT_TEXT, size, bold=bold)
        return p

    title_line("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", size=12)
    title_line("РОССИЙСКОЙ ФЕДЕРАЦИИ", size=12)
    title_line("ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ", size=12, before=6)
    title_line("ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", size=12)
    title_line("«МОСКОВСКИЙ АВИАЦИОННЫЙ ИНСТИТУТ", size=12, bold=True, before=6)
    title_line("(НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ)»", size=12, bold=True)
    title_line("Институт № 8 «Компьютерные науки и прикладная математика»",
               size=12, before=6)
    title_line("Направление: фундаментальная информатика", size=12, before=6)
    title_line("и информационные технологии", size=12, after=60)

    title_line("КУРСОВАЯ РАБОТА (ПРОЕКТ)", bold=True, size=16, after=6)
    title_line("по системному программированию", size=14, after=48)

    title_line("на тему:", size=14, after=6)
    title_line("«Система управления базами данных,", size=14, bold=True)
    title_line("вариант 1 (B-tree)»", size=14, bold=True, after=60)

    def signature_line(text, before=0, after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(9)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        set_run_font(run, FONT_TEXT, 12)
        return p

    signature_line("Выполнил:")
    signature_line("студент группы М8О-214БВ-24")
    signature_line("Богданов М. А.                            ____________", after=12)
    signature_line("Проверил:")
    signature_line("_______________________________")
    signature_line("_______________________________   ____________")

    title_line("Москва, 2026", size=12, before=120)

    doc.add_page_break()

    # ── СОДЕРЖАНИЕ ────────────────────────────────────────────────────
    add_heading(doc, "Содержание", level=1)
    add_text(doc,
        "Оглавление формируется автоматически в Word: вкладка «Ссылки» "
        "→ «Оглавление» → выбрать стиль. Все Heading 1 и Heading 2 "
        "разделов попадут в оглавление со сквозной нумерацией.")
    doc.add_page_break()

    # ── ВВЕДЕНИЕ ──────────────────────────────────────────────────────
    build_intro(doc)
    doc.add_page_break()

    # ── ОСНОВНАЯ ЧАСТЬ ────────────────────────────────────────────────
    add_heading(doc, "Основная часть", level=1)
    build_point0(doc)
    doc.add_page_break()
    build_addon_grpc(doc)
    doc.add_page_break()
    build_addon_logs(doc)
    doc.add_page_break()
    build_addon_telemetry(doc)
    doc.add_page_break()
    build_addon_default(doc)
    doc.add_page_break()
    build_addon_andor(doc)
    doc.add_page_break()
    build_addon_aggregates(doc)
    doc.add_page_break()

    # ── ВЫВОД ─────────────────────────────────────────────────────────
    build_conclusion(doc)
    doc.add_page_break()

    # ── ИСТОЧНИКИ ─────────────────────────────────────────────────────
    build_sources(doc)
    doc.add_page_break()

    # ── ПРИЛОЖЕНИЯ ────────────────────────────────────────────────────
    build_appendices(doc)

    # ── СОХРАНЕНИЕ ────────────────────────────────────────────────────
    out_path = "/home/study/coursework/theory/zapiska/poyasnitelnaya_zapiska.docx"
    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    build()
